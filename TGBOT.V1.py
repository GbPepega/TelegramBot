import telebot
import datetime
from telebot import types
from docx import Document


bot = telebot.TeleBot('xxxxxxxxxxx')
output_folder = "xxxxxxxxxxxxxx"


states = {}

main_menu_keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
signup_button = types.KeyboardButton("Записаться")
contacts_button = types.KeyboardButton("Контакты")
main_menu_keyboard.row(contacts_button, signup_button)

schedule_keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
days = ["Понедельник 19:00", "Среда 19:00", "Пятница 19:00"]
day_buttons = [types.KeyboardButton(day) for day in days]
schedule_keyboard.row(*day_buttons)

info_menu_keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
info_types = ["ФИО", "Номер телефона", "Цель посещения"]
info_buttons = [types.KeyboardButton(info) for info in info_types]
send_info_button = types.KeyboardButton("Отправить информацию")
info_menu_keyboard.row(*info_buttons)
info_menu_keyboard.row(send_info_button)

def send_error_message(chat_id, text="Произошла ошибка при обработке вашего запроса. Пожалуйста, попробуйте еще раз."):
    bot.send_message(chat_id, text)

@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(message.chat.id, "Добро пожаловать! Я бот филиала стрелкового общества Евразия. Выберите действие:", reply_markup=main_menu_keyboard)
    states[message.chat.id] = "main_menu"

@bot.message_handler(func=lambda message: message.text == "Записаться")
def handle_signup(message):
    bot.send_message(message.chat.id, "Выберите день недели:", reply_markup=schedule_keyboard)
    states[message.chat.id] = "schedule_menu"

@bot.message_handler(func=lambda message: message.text == "Контакты")
def handle_contacts(message):
    contact_info = (
        "Наши контакты:\n"
        "Телефон: +1 234 567 890\n"
        "Email: example@email.com\n"
        "Адрес: Город, Улица, Дом\n"
        "Instagram: @example_instagram\n"
        "Дополнительная информация: что-то ещё\n"
    )
    bot.send_message(message.chat.id, contact_info)


@bot.message_handler(func=lambda message: message.text in days)
def handle_schedule_day(message):
    day = message.text
    bot.send_message(message.chat.id, f"Вы выбрали {day}. Теперь выберите информацию:", reply_markup=info_menu_keyboard)
    states[message.chat.id] = {"schedule_day": day, "info_menu": {}}

@bot.message_handler(func=lambda message: message.text in info_types)
def handle_info(message):
    info_type = message.text
    bot.send_message(message.chat.id, f"Вы выбрали {info_type}. Теперь введите информацию:")
    states[message.chat.id]["current_info"] = info_type
    states[message.chat.id]["current_input"] = ""

@bot.message_handler(func=lambda message: message.text and states.get(message.chat.id, {}).get("current_input") is not None)
def handle_input(message):
    chat_id = message.chat.id
    current_info = states[chat_id]["current_info"]
    current_input = message.text
    
    if current_info is not None:
        # Добавляем проверку на отправку ссылок и фотографий
        if 'http' in current_input or message.content_type == 'photo':
            send_error_message(chat_id, "Пожалуйста, введите корректную информацию без ссылок и фото")
            return

        states[chat_id]["info_menu"][current_info] = current_input
        bot.send_message(chat_id, f"Вы ввели следующую информацию для пункта {current_info}: {current_input}")
        states[chat_id]["current_input"] = None
    else:
        send_error_message(chat_id)

@bot.message_handler(func=lambda message: message.text == "Отправить информацию")
def handle_send_info(message):
    chat_id = message.chat.id
    current_time = datetime.datetime.now()
    schedule_day = states.get(chat_id, {}).get("schedule_day")
    info_menu = states.get(chat_id, {}).get("info_menu", {})
    info_sent = states.get(chat_id, {}).get("info_sent", False)


  

    if info_sent:
        bot.send_message(chat_id, "Заявка уже отправлена. Если у вас есть еще вопросы, обратитесь к администратору. Пожалуйста, перезапустите бот, чтобы отправить новую информацию.")
        return

    if all([schedule_day, info_menu, all(info_menu.values())]):
        full_info = {"День недели": schedule_day, **info_menu}

        doc = Document()
        doc.add_heading('Информация о пользователе', level=1)

        for key, value in full_info.items():
            doc.add_paragraph(f"{key}: {value}")

        filename = f"{current_time}пользователь_с_id_{chat_id}_info.docx"
        file_path = f"{output_folder}/{filename}"

        try:
            doc.save(file_path)
        except Exception as e:
            send_error_message(chat_id, f"Ошибка при сохранении: {e}")
            return

        bot.send_message(chat_id, "Информация отправлена, вы записаны (Также можно связаться по телефону или WhatsApp, номера указаны в описании)")

        states[chat_id]["info_sent"] = True
        states[chat_id]["schedule_day"] = None
        states[chat_id]["info_menu"] = {}

    else:
        bot.send_message(chat_id, "Вы не заполнили все поля. Пожалуйста, введите всю недостающую информацию")

if __name__ == "__main__":
    bot.polling(none_stop=True)
